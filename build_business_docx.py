from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "AI问诊电子病历系统商业化技术文档.md"
OUTPUT = ROOT / "AI问诊电子病历系统商业化技术文档.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "AI 问诊电子病历生成系统商业化技术文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    add_field(footer, "PAGE")
    footer.add_run(" 页")
    for r in footer.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(100, 100, 100)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    title = p.add_run("AI 问诊电子病历生成系统")
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = p2.add_run("商业化技术文档")
    subtitle.font.size = Pt(18)
    subtitle.font.color.rgb = RGBColor.from_string("2E74B5")
    subtitle.font.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(36)
    meta.add_run("版本：V1.0\n日期：2026-06-08\n适用范围：门诊问诊、基层医疗、专科门诊、互联网医院、体检中心、诊所集团")

    doc.add_section(WD_SECTION.NEW_PAGE)


def parse_inline(text):
    return re.sub(r"`([^`]+)`", r"\1", text).strip()


def add_markdown_paragraph(doc, line):
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("- "):
        p = doc.add_paragraph(parse_inline(stripped[2:]), style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        return
    if re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(parse_inline(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        return
    p = doc.add_paragraph(parse_inline(stripped))
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, rows):
    parsed_rows = []
    for row in rows:
        cells = [parse_inline(cell.strip()) for cell in row.strip().strip("|").split("|")]
        parsed_rows.append(cells)
    if len(parsed_rows) < 2:
        return
    if all(re.fullmatch(r"\s*:?-{3,}:?\s*", c) for c in parsed_rows[1]):
        data = [parsed_rows[0]] + parsed_rows[2:]
    else:
        data = parsed_rows
    col_count = max(len(row) for row in data)
    table = doc.add_table(rows=len(data), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(data):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = text
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
    widths = table_widths(col_count)
    set_table_width(table, widths)
    doc.add_paragraph()


def table_widths(col_count):
    total = 9360
    if col_count == 2:
        return [2600, 6760]
    if col_count == 3:
        return [2200, 3180, 3980]
    if col_count == 4:
        return [1800, 2400, 2800, 2360]
    return [total // col_count] * col_count


def build_docx():
    doc = Document()
    style_document(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_buffer = []
    code_buffer = []
    in_code = False
    skip_title = True

    for line in lines:
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_buffer))
                p.style = "No Spacing"
                p.paragraph_format.left_indent = Inches(0.25)
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                flush_table(doc, table_buffer)
                table_buffer.clear()
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            continue
        flush_table(doc, table_buffer)
        table_buffer.clear()

        if line.startswith("# "):
            if skip_title:
                skip_title = False
                continue
            doc.add_heading(parse_inline(line[2:]), level=1)
        elif line.startswith("## "):
            doc.add_heading(parse_inline(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(parse_inline(line[4:]), level=2)
        elif line.startswith("#### "):
            doc.add_heading(parse_inline(line[5:]), level=3)
        else:
            add_markdown_paragraph(doc, line)

    flush_table(doc, table_buffer)
    doc.save(OUTPUT)


def flush_table(doc, table_buffer):
    if table_buffer:
        add_table(doc, table_buffer)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
