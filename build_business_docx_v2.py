from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "AI问诊电子病历系统商业化技术文档v2.0.docx"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MID_BLUE = "1F4D78"
LIGHT_FILL = "F2F4F7"
SOFT_BLUE = "E8EEF5"
GREEN = "1F7A4D"
AMBER = "7A5A00"
RED = "9B1C1C"
GRAY = "666666"


def set_east_asia(run, font="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_field(paragraph, instr: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(text)
    run._r.append(end)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, bottom=90, start=130, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def style_table(table, widths, header_fill=LIGHT_FILL):
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    set_east_asia(run)
                    run.font.size = Pt(9.3)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths is None:
        widths = [int(9360 / len(headers))] * len(headers)
    style_table(table, widths, header_fill)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text="", bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_east_asia(r1)
        r1.font.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        set_east_asia(r2)
    else:
        r = p.add_run(text)
        set_east_asia(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_east_asia(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_east_asia(r)


def add_callout(doc, label, text, fill=SOFT_BLUE, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + "：")
    set_east_asia(r1)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor.from_string(color)
    r2 = p.add_run(text)
    set_east_asia(r2)
    r2.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_table_geometry(table, [9360])
    doc.add_paragraph()


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, MID_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("AI问诊电子病历系统商业化技术文档 v2.0")
    set_east_asia(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_east_asia(run)
    add_field(footer, "PAGE")
    run = footer.add_run(" 页")
    set_east_asia(run)
    for r in footer.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(GRAY)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(115)
    title = p.add_run("AI问诊电子病历系统")
    set_east_asia(title)
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = p2.add_run("商业化技术文档")
    set_east_asia(sub)
    sub.font.size = Pt(19)
    sub.font.bold = True
    sub.font.color.rgb = RGBColor.from_string(BLUE)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(34)
    meta = p3.add_run("版本：v2.0\n日期：2026-06-09\n适用范围：门诊问诊、基层医疗、诊所集团、互联网医院、体检中心、专科门诊")
    set_east_asia(meta)
    meta.font.size = Pt(11)
    meta.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    add_callout(
        doc,
        "文档定位",
        "本文件面向商业化落地，描述当前系统能力、技术架构、部署方案、安全合规边界、产品化路线和可规模化优化方向，可作为项目立项、合作洽谈、研发迭代和试点上线的基础技术资料。",
        fill="EEF5FF",
    )
    doc.add_section(WD_SECTION.NEW_PAGE)


def build():
    doc = Document()
    setup_doc(doc)
    add_cover(doc)

    add_heading(doc, "1. 执行摘要", 1)
    add_para(doc, "AI问诊电子病历系统 v2.0 是一套面向门诊场景的智能病历生成平台，核心流程为“登录系统、上传录音或实时录音、语音转写、AI病历草稿、医生确认、打印和导出”。系统当前已完成可演示、可小范围试点的闭环，并具备向商业化 SaaS 或私有化部署演进的基础。")
    add_bullets(doc, [
        "实时语音识别：已从 Vosk 小模型升级为 sherpa-onnx streaming Paraformer 中文流式模型，支持浏览器 WebSocket 边录边转写。",
        "AI病历生成：已接入 DeepSeek 大模型，支持从医患对话中提取门诊通用病历字段，并生成候选诊断供医生选择和编辑。",
        "医生确认闭环：医生可修改主诉、现病史、既往史、过敏史、体格检查、初步诊断和处理意见后确认保存。",
        "商业化基础：已具备登录、医生姓名展示、审计日志、任务队列雏形、文件清理、HIS/EMR集成桩、打印和 TXT/JSON 导出能力。",
        "当前边界：SQLite、单机进程、简化权限、基础审计和临时公网隧道仍不满足正式大规模上线，需要按本文件路线升级。"
    ])
    add_callout(doc, "商业判断", "当前系统适合用于内部测试、小范围试点、产品演示和需求验证；若面向多机构正式收费使用，需完成数据库、权限、安全、审计、任务队列、对象存储、监控和合规治理升级。", fill="FFF8E5", color=AMBER)

    add_heading(doc, "2. 产品目标与适用场景", 1)
    add_para(doc, "系统目标不是替代医生，而是降低门诊病历书写负担，将自然医患对话转化为结构化、可确认、可打印、可导出的电子病历草稿。所有诊断和处置内容必须由医生最终确认。")
    add_table(doc, ["场景", "用户", "核心价值", "商业化机会"], [
        ["基层门诊", "全科医生、诊所医生", "减少重复书写，提高接诊效率", "按医生账号或门诊量订阅"],
        ["专科门诊", "内科、耳鼻喉、皮肤、康复等", "沉淀专科模板和术语纠错", "专科包、模板包增值"],
        ["互联网医院", "线上问诊医生", "将语音/文本问诊快速转为病历", "API 调用量计费"],
        ["体检与慢病随访", "健康管理师、医生", "生成随访记录和干预建议草稿", "机构私有化部署"],
        ["诊所集团", "多门店运营方", "统一病历规范和审计追踪", "集团版 SaaS 或本地化交付"],
    ], widths=[1500, 1800, 3300, 2760])

    add_heading(doc, "3. 当前系统能力清单", 1)
    add_table(doc, ["模块", "当前能力", "商业化状态", "后续增强"], [
        ["登录与医生身份", "登录后进入工作台，页面显示当前医生姓名", "可试点", "多角色、多机构、科室权限"],
        ["音频上传", "支持上传录音文件并触发转写", "可试点", "大文件断点、批量任务、对象存储"],
        ["实时录音转写", "浏览器麦克风采集，WebSocket 推送 16k PCM，后端实时识别", "已升级", "说话人分离、标点恢复、医学纠错"],
        ["ASR 引擎", "默认 sherpa-onnx streaming Paraformer；保留 Vosk 与 Whisper 相关能力", "可试点", "模型热切换、GPU/CPU 自适应"],
        ["AI病历草稿", "DeepSeek 生成门诊通用病历字段和候选诊断", "可试点", "专科模板、质控规则、引用证据"],
        ["医生确认", "医生编辑后确认保存", "可试点", "电子签名、修改留痕、版本对比"],
        ["打印导出", "支持打印病历、导出 TXT 和 JSON", "可试点", "PDF、DOCX、院内 EMR 格式"],
        ["审计日志", "记录登录、上传、转写、生成、确认、导出等动作", "基础版", "审计检索、异常告警、留痕报表"],
        ["集成接口", "已有 HIS/EMR 同步桩和集成日志", "原型", "标准 HL7/FHIR/院内接口适配"],
    ], widths=[1500, 3000, 1500, 3360])

    add_heading(doc, "4. 目标业务流程", 1)
    add_numbers(doc, [
        "医生登录系统，进入工作台，系统展示医生姓名和基础操作区。",
        "医生可选择上传既有录音，或点击开始实时录音。",
        "系统将语音转为文本；实时模式下，浏览器通过 WebSocket 连续推送音频流，后端持续返回 partial/final 文本。",
        "医生确认或补充转写文本后，点击生成草稿。",
        "DeepSeek 根据对话内容输出结构化门诊病历草稿，并列出可能疾病候选项、依据和建议检查。",
        "医生选择或手动编辑初步诊断，完善体格检查、处理意见等字段。",
        "医生确认病历，系统保存结构化记录、原始转写、AI草稿和审计日志。",
        "系统支持打印门诊病历，或导出 TXT/JSON 供归档、接口同步或后续集成使用。"
    ])

    add_heading(doc, "5. 技术架构设计", 1)
    add_para(doc, "当前 v2.0 采用轻量单体架构，前端为原生 HTML/CSS/JavaScript，后端为 FastAPI，数据层为 SQLite，AI能力由本地 ASR 模型和 DeepSeek 大模型共同完成。商业化正式上线建议演进为“Web/API/任务队列/模型服务/数据库/对象存储/监控审计”分层架构。")
    add_table(doc, ["层级", "当前实现", "职责", "商业化演进"], [
        ["前端工作台", "index.html、app.js、styles.css", "录音、上传、转写展示、草稿编辑、打印导出", "组件化前端、权限路由、移动端适配"],
        ["API 服务", "FastAPI backend/main.py", "认证、文件上传、转写任务、病历生成、审计", "多进程部署、OpenAPI 管理、限流"],
        ["实时通道", "/ws/asr/stream、/ws/jobs/{job_id}", "实时 ASR 和任务状态推送", "鉴权 WebSocket、断线续传、房间隔离"],
        ["ASR 模型服务", "sherpa-onnx Paraformer 本地推理", "低延迟中文语音识别", "独立模型服务、GPU池、模型版本管理"],
        ["LLM 服务", "DeepSeek OpenAI 兼容接口", "病历结构化、候选诊断、医学表达优化", "供应商路由、缓存、质控、多模型评测"],
        ["数据层", "SQLite 文件数据库", "音频记录、病历、审计、集成日志", "PostgreSQL/MySQL、读写分离、备份恢复"],
        ["文件层", "data/uploads、data/exports", "录音、分片、导出文件", "对象存储、生命周期、加密、病毒扫描"],
        ["集成层", "EMR 同步桩", "预留院内系统对接", "HL7/FHIR/厂商接口适配器"],
    ], widths=[1450, 2200, 2600, 3110])

    add_heading(doc, "6. AI语音识别方案 v2.0", 1)
    add_para(doc, "实时转写模块是 v2.0 的关键升级。此前 Vosk 小模型在中文口语和医疗对话中出现误识别、重复和漏识别问题；当前已升级为 sherpa-onnx streaming Paraformer bilingual zh-en 模型，采用 int8 ONNX 模型在 CPU 上实时推理。")
    add_table(doc, ["能力项", "v1/Vosk 小模型", "v2/sherpa-onnx Paraformer", "商业化建议"], [
        ["实时性", "可流式，但准确率有限", "低延迟 partial 输出，测试音频约 3.7 秒端到端完成", "独立 ASR 服务，按并发扩容"],
        ["准确率", "普通话短句可用，医疗对话不稳定", "中文连续对话更完整，口语识别明显改善", "加入医学术语纠错和热词增强"],
        ["部署方式", "本地模型", "本地 ONNX int8 模型，CPU 可运行", "按机构数据安全要求私有化部署"],
        ["方言支持", "有限", "模型文档支持普通话及部分方言能力", "结合方言模型或普通话归一化层"],
        ["成本", "低", "低，主要消耗本机 CPU", "高并发场景评估 CPU/GPU 资源池"],
    ], widths=[1500, 2300, 3100, 2460])
    add_bullets(doc, [
        "浏览器采集麦克风音频，并在前端降采样为 16kHz 16bit PCM。",
        "前端通过 WebSocket 将 PCM 音频片段发送到 /ws/asr/stream。",
        "后端创建 sherpa-onnx OnlineRecognizer 和 stream，连续 accept_waveform 并 decode_stream。",
        "识别过程中向前端返回 partial 文本，停止录音后返回 final 文本和 done 事件。",
        "文本进入前端转写框，医生可直接编辑，再提交给 DeepSeek 生成病历草稿。"
    ])
    add_callout(doc, "优化方向", "ASR 不应只追求模型替换，还需要热词、标点、说话人角色、医学术语纠错和噪声处理形成组合能力。医疗商业化场景中，ASR 后处理质量常常决定最终病历可用度。", fill="EEF5FF")

    add_heading(doc, "7. DeepSeek病历生成方案", 1)
    add_para(doc, "系统当前通过 DeepSeek OpenAI-compatible Chat Completions JSON 模式生成门诊通用病历。提示词要求模型基于医患对话输出结构化字段，并使用专业医学术语表达。对于诊断，系统不再通过前端 if 规则硬判，而是要求大模型列出候选疾病选项，医生可点击选择或手动编辑。")
    add_table(doc, ["输出字段", "说明", "医生确认要求"], [
        ["主诉", "提炼患者主要症状及持续时间", "必须核对症状与时间"],
        ["现病史", "按发生、发展、伴随症状、诱因、缓解因素整理", "需补充阴性症状和重要时间线"],
        ["既往史", "提取高血压、糖尿病、手术史等", "患者未提及时不得臆造"],
        ["过敏史", "提取药物、食物及其他过敏信息", "缺失时标注未提及或待补充"],
        ["体格检查", "根据对话生成待查或已查内容", "未实际检查不得写成已完成事实"],
        ["初步诊断", "生成可编辑初步诊断", "必须由医生最终确认"],
        ["候选诊断", "列出 3-6 个可能疾病、依据和建议检查", "供医生选择，不直接替代诊断"],
        ["处理意见", "生成检查、治疗、随访和健康教育建议草稿", "需结合本机构规范和医嘱权限"],
        ["风险提醒", "提示红旗症状、信息不足和需转诊风险", "作为质控参考"],
    ], widths=[1600, 4300, 3460])

    add_heading(doc, "8. 数据模型与文件管理", 1)
    add_para(doc, "当前数据库为 SQLite，主要表包括 audio_files、emr_records、audit_logs、integration_logs。此设计便于本地演示和小范围试点，但不适合多用户高并发生产环境。")
    add_table(doc, ["数据对象", "当前存储", "用途", "商业化改造"], [
        ["音频文件", "data/uploads + audio_files", "存储上传文件和元信息", "对象存储、加密、生命周期策略"],
        ["病历记录", "emr_records", "保存患者信息、转写文本、病历 JSON、证据和确认时间", "版本化、电子签名、归档策略"],
        ["审计日志", "audit_logs", "记录关键用户行为", "完整审计页面、检索、导出、异常检测"],
        ["集成日志", "integration_logs", "记录 EMR/HIS 同步请求与结果", "重试队列、幂等键、失败告警"],
        ["导出文件", "data/exports", "存放导出内容", "短期缓存、权限校验、到期清理"],
    ], widths=[1600, 2100, 2900, 2760])
    add_bullets(doc, [
        "正式商业化不建议继续使用 SQLite，应迁移至 PostgreSQL 或 MySQL。",
        "录音文件属于敏感健康数据，应设置最短必要保存周期，并支持按机构策略自动删除。",
        "审计日志不应出现在打印病历内，应作为后台管理和合规追溯能力独立存在。",
        "病历确认后应保存 AI 原始草稿、医生最终确认版本、修改时间和操作人，便于责任边界清晰。"
    ])

    add_heading(doc, "9. 安全、权限与合规设计", 1)
    add_para(doc, "医疗数据涉及个人敏感信息和健康信息。商业化版本必须将安全设计从“功能附属”提升为“平台底座”。当前登录功能可用于试点，但正式上线需要机构、角色、科室、医生、管理员、审计员等多层权限模型。")
    add_table(doc, ["安全域", "当前状态", "正式上线要求"], [
        ["身份认证", "单管理员账号和 Cookie 会话", "机构多租户、强密码、MFA、会话过期、单点登录"],
        ["授权模型", "基础登录保护", "RBAC/ABAC，按机构、科室、医生、病历范围隔离"],
        ["传输安全", "本地 HTTP 和临时 Cloudflare tunnel", "HTTPS、WSS、正式域名和证书"],
        ["数据安全", "本地文件和 SQLite", "数据库加密、对象存储加密、密钥管理、最小权限"],
        ["隐私合规", "原型阶段", "用户授权、告知同意、数据最小化、删除和导出机制"],
        ["审计追踪", "基础日志", "不可篡改日志、审计检索、异常行为告警"],
        ["模型安全", "直接调用 DeepSeek", "脱敏策略、提示词安全、输出质控、供应商 SLA"],
    ], widths=[1600, 3200, 4560])

    add_heading(doc, "10. 部署架构与上线方案", 1)
    add_para(doc, "系统可按三种形态商业化：演示/试点版、机构私有化版、云端 SaaS 版。当前代码更接近演示/试点版；若要给外部客户长期使用，应至少完成生产级部署改造。")
    add_table(doc, ["部署形态", "适用客户", "优势", "限制"], [
        ["本地试点版", "单诊所、内部测试", "部署快、成本低、便于验证流程", "安全、并发、运维能力有限"],
        ["私有化部署", "医院、连锁诊所、数据敏感机构", "数据留在客户环境，易对接院内系统", "交付和运维成本较高"],
        ["云端 SaaS", "中小诊所、轻量门诊", "快速开通、统一升级、可规模化收费", "需严格多租户隔离和合规治理"],
        ["混合部署", "大型医疗集团", "本地 ASR/数据 + 云端 LLM 或管理平台", "架构复杂，需要清晰责任边界"],
    ], widths=[1600, 2200, 3000, 2560])
    add_numbers(doc, [
        "将 FastAPI 服务迁移到生产 ASGI 部署，如 Gunicorn/Uvicorn workers 或容器化编排。",
        "将 SQLite 迁移为 PostgreSQL/MySQL，并设计迁移脚本和备份策略。",
        "将 ASR 推理拆分为独立模型服务，避免模型推理阻塞 API 主进程。",
        "引入 Redis/RQ/Celery 等任务队列，承载长音频转写、批量导入、导出和同步任务。",
        "将上传文件迁移到对象存储，并设置访问签名、病毒扫描、生命周期删除策略。",
        "配置 HTTPS/WSS、正式域名、反向代理、访问日志、限流和 WAF。",
        "接入监控告警：API 延迟、转写耗时、LLM失败率、任务堆积、磁盘空间、模型内存占用。"
    ])

    add_heading(doc, "11. 性能与准确率优化路线", 1)
    add_para(doc, "系统性能由三部分共同决定：前端采集与传输、ASR 模型推理、LLM 生成与后处理。商业化版本应建立可量化评测集，而不是只凭主观听感判断。")
    add_table(doc, ["优化方向", "具体措施", "预期收益", "优先级"], [
        ["ASR 热词", "加入科室疾病、药品、检查、常见症状词表", "减少同音字和医学术语错误", "高"],
        ["医学纠错层", "DeepSeek 对转写文本做只纠错不改意的普通话和医学术语归一化", "提高病历生成质量", "高"],
        ["说话人角色", "识别医生/患者轮次，或通过 UI 快捷标注", "改善现病史结构化", "高"],
        ["标点恢复", "流式输出后做句读和停顿修复", "提高可读性", "中"],
        ["噪声处理", "前端增益控制、降噪、静音检测", "减少漏识别和误触发", "中"],
        ["模型服务化", "ASR 独立进程池，按并发扩展", "避免 API 阻塞", "高"],
        ["LLM 缓存与压缩", "长对话摘要、分段抽取、结果缓存", "降低生成耗时和成本", "中"],
        ["评测体系", "建立 WER/CER、字段抽取准确率、医生修改率指标", "支撑商业验收", "高"],
    ], widths=[1700, 3900, 2500, 1260])

    add_heading(doc, "12. HIS/EMR 集成设计", 1)
    add_para(doc, "当前系统已有 EMR 同步接口桩和 integration_logs，说明架构上已预留集成能力。商业化落地时，集成层需要独立设计为可配置适配器，避免将不同医院接口逻辑写死在主业务代码中。")
    add_bullets(doc, [
        "接口方向：病人基本信息同步、挂号/就诊号同步、病历回写、检查检验申请、处方或医嘱回写、打印模板对齐。",
        "标准协议：优先评估 HL7、FHIR、WebService、REST、数据库视图、中间库、院内厂商私有接口。",
        "可靠性：同步接口应支持幂等键、失败重试、死信队列、人工重发和完整集成日志。",
        "安全性：院内接口凭证应独立管理，不写入前端，不进入普通日志。",
        "责任边界：AI 生成内容必须在医生确认后再回写 EMR，未确认草稿不得作为正式病历。"
    ])

    add_heading(doc, "13. 商业化版本规划", 1)
    add_table(doc, ["版本", "目标客户", "核心能力", "收费方式"], [
        ["试点版", "单机构、内部验证", "登录、录音转写、AI草稿、医生确认、打印导出", "项目试点费或免费 PoC"],
        ["标准版", "中小诊所、门诊部", "多医生账号、审计、模板、数据备份、基础统计", "按医生/月订阅"],
        ["专业版", "专科门诊、连锁诊所", "专科模板、医学纠错、质控规则、批量任务、API集成", "账号费 + 用量费"],
        ["私有化版", "医院、集团客户", "本地部署、HIS/EMR对接、专属模型配置、运维监控", "一次性交付 + 年维护费"],
        ["平台版", "区域或集团平台", "多租户、多机构、统一监控、数据治理、模型评测", "平台授权 + 服务费"],
    ], widths=[1400, 2000, 4300, 1660])

    add_heading(doc, "14. 运维监控与质量指标", 1)
    add_para(doc, "商业化系统需要将“能用”升级为“可观测、可恢复、可追责”。建议至少建立以下指标体系。")
    add_table(doc, ["指标类别", "关键指标", "说明"], [
        ["可用性", "服务可用率、接口错误率、WebSocket断开率", "衡量系统稳定性"],
        ["性能", "实时转写延迟、LLM生成耗时、上传转写耗时", "衡量用户等待时间"],
        ["准确率", "ASR字错误率、医学术语错误率、病历字段抽取准确率", "衡量AI质量"],
        ["业务", "日活医生数、生成病历数、医生平均修改率、确认率", "衡量商业价值"],
        ["成本", "单次病历 LLM 成本、ASR CPU/GPU成本、存储成本", "支撑定价模型"],
        ["安全", "登录失败次数、异常导出、越权访问、审计查询", "支撑合规和风控"],
    ], widths=[1600, 3600, 4160])

    add_heading(doc, "15. 风险分析与控制措施", 1)
    add_table(doc, ["风险", "表现", "影响", "控制措施"], [
        ["ASR误识别", "同音字、漏句、方言误识别", "病历内容错误", "医生确认、医学纠错、热词、评测集"],
        ["LLM幻觉", "生成未提及病史或诊断依据", "医疗风险", "提示词约束、证据引用、缺失标注、医生确认"],
        ["数据泄露", "录音、病历、密钥泄露", "合规和品牌风险", "加密、权限、审计、密钥管理"],
        ["并发不足", "多人同时转写卡顿", "试点体验差", "模型服务化、队列、资源池"],
        ["集成失败", "回写 EMR 失败或重复", "业务中断", "幂等、重试、集成日志、人工补偿"],
        ["责任边界不清", "AI草稿被误认为正式诊断", "医疗责任风险", "明确“医生最终确认”流程和页面提示"],
    ], widths=[1600, 2600, 2100, 3060])

    add_heading(doc, "16. 12周商业化落地路线图", 1)
    add_table(doc, ["阶段", "周期", "目标", "交付物"], [
        ["阶段一：试点稳定", "第1-2周", "修复实时转写体验和病历字段质量", "ASR评测集、医学纠错、错误日志"],
        ["阶段二：权限与数据", "第3-4周", "多账号、多机构、数据库迁移", "PostgreSQL、RBAC、迁移脚本"],
        ["阶段三：生产部署", "第5-6周", "容器化、HTTPS、对象存储、任务队列", "Docker Compose/K8s方案、队列服务"],
        ["阶段四：审计与合规", "第7-8周", "完整审计页面、数据生命周期、操作留痕", "审计后台、清理策略、导出记录"],
        ["阶段五：集成与模板", "第9-10周", "对接试点机构 HIS/EMR 和专科模板", "接口适配器、模板中心"],
        ["阶段六：商业验收", "第11-12周", "性能、准确率、可用性和医生满意度验收", "验收报告、SLA、上线手册"],
    ], widths=[1800, 1300, 3300, 2960])

    add_heading(doc, "17. 结论", 1)
    add_para(doc, "AI问诊电子病历系统 v2.0 已形成从语音到病历草稿再到医生确认和打印导出的核心闭环，并完成实时转写模型升级和 DeepSeek 病历生成集成。系统已经具备小范围试点和商业演示价值，但距离正式规模化上线仍需要生产级数据库、权限体系、队列、模型服务化、完整审计、文件生命周期、HIS/EMR集成和合规治理。")
    add_para(doc, "建议下一阶段以“准确率、稳定性、合规性、可集成性”为主线推进，而不是继续堆叠界面功能。只有当医生修改率下降、生成耗时稳定、审计可追溯、数据安全可证明时，系统才真正具备持续收费和规模复制的商业基础。")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
